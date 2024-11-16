import pandas as pd
from docx import Document
import re

# сейчас парсятся все таблицы из 1 документа. самое главное сейчас - чтобы хэдэр
# состоял из 2х строк объединенных. сейчас делает только первую
# Функция для извлечения данных из таблицы, объединяя разорванные строки и обрабатывая тире
def extract_table_data_with_merged_cells(table):
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # Удаляем лишние пробелы и обрабатываем тире
            text = re.sub(r"\s+", " ", cell.text.strip().replace("\n", " "))
            text = "0" if text == "–" or text == "-" else text
            row_data.append(text)
        table_data.append(row_data)
    return table_data


# Функция для "разворачивания" объединённых ячеек
def fill_merged_cells(data):
    filled_data = []
    for row in data:
        filled_row = []
        for i, cell in enumerate(row):
            if not cell and filled_data:
                cell = filled_data[-1][i]
            filled_row.append(cell)
        filled_data.append(filled_row)
    return filled_data


# Функция для очистки и проверки уникальности заголовков
def clean_headers(headers):
    cleaned_headers = []
    header_count = {}
    for header in headers:
        if not header:  # Если заголовок пустой
            header = "Unnamed"
        if header in header_count:  # Если заголовок уже существует
            header_count[header] += 1
            header = f"{header}_{header_count[header]}"  # Добавляем суффикс для уникальности
        else:
            header_count[header] = 0
        cleaned_headers.append(header)
    return cleaned_headers


# Функция для проверки уникальности заголовков и их обновления
def ensure_unique_columns(dataframe):
    columns = list(dataframe.columns)
    unique_columns = clean_headers(columns)  # Применяем функцию clean_headers для уникальности
    dataframe.columns = unique_columns  # Обновляем столбцы DataFrame
    return dataframe


# Функция для проверки, является ли таблица сноской
def is_footnote_table(header):
    if re.match(r"^\d+\)", header) or re.match(r"^\d+\.\)", header):
        return True
    return False


# Функция для выравнивания количества столбцов
def align_columns(data, expected_columns_count):
    aligned_data = []
    for row in data:
        if len(row) < expected_columns_count:
            # Дополняем недостающие столбцы пустыми значениями
            row += [''] * (expected_columns_count - len(row))
        elif len(row) > expected_columns_count:
            # Обрезаем лишние столбцы
            row = row[:expected_columns_count]
        aligned_data.append(row)
    return aligned_data


# Функция для проверки, совпадают ли заголовки двух таблиц
def headers_match(table1, table2):
    if not table1 or not table2:
        return False
    header1 = clean_headers(table1[0])
    header2 = clean_headers(table2[0])
    return header1 == header2


# Функция для проверки, являются ли таблицы частью одной структуры с регионами слева и справа
def is_combined_table(left_table_data, right_table_data):
    if not left_table_data or not right_table_data:
        return False
    left_regions = [row[0] for row in left_table_data if "область" in row[0] or "республика" in row[0]]
    right_regions = [row[-1] for row in right_table_data if "область" in row[-1] or "республика" in row[-1]]
    return bool(left_regions) and bool(right_regions)


# Функция для удаления строк, где в первом столбце нет данных о регионе (кроме первых двух строк)
def remove_empty_region_rows(df):
    if df.empty:
        return df

    # Предполагаем, что регион/область/край находится в первом столбце
    region_column = df.columns[0]

    # Оставляем только те строки, где первый столбец содержит значимые данные
    # Пропускаем только первые две строки
    df = df[(df[region_column].notnull()) | (df.index < 2)]

    return df


# Открытие документа Word
doc_path = "/content/R_01.docx"  # Замените на путь к вашему файлу
doc = Document(doc_path)

# Переменные для хранения объединённых таблиц
merged_tables = {}
num_tables = len(doc.tables)
i = 0

# Переменная для хранения "текущей" таблицы
current_df = None

while i < num_tables:
    try:
        # Извлекаем данные из текущей таблицы
        left_table_data = fill_merged_cells(extract_table_data_with_merged_cells(doc.tables[i]))

        # Пропускаем таблицы-сноски
        if left_table_data:
            left_header_candidate = left_table_data[0][0] if left_table_data[0] else ""
            if is_footnote_table(left_header_candidate):
                print(f"Пропуск таблицы {i + 1} из-за сноски.")
                i += 1
                continue

        # Проверяем, если текущая таблица — продолжение предыдущей
        if current_df is not None:
            # Проверяем, совпадают ли заголовки текущей таблицы и текущего DataFrame
            current_headers = list(current_df.columns)
            left_headers = clean_headers(left_table_data[0]) if left_table_data else []

            if current_headers == left_headers:
                print(f"Таблица {i + 1} считается продолжением предыдущей таблицы.")
                # Добавляем строки в текущую таблицу
                left_table_data = align_columns(left_table_data[1:], len(current_headers))
                df_part = pd.DataFrame(left_table_data, columns=current_headers)
                current_df = pd.concat([current_df, df_part], ignore_index=True)
                i += 1
                continue

                # Обрабатываем новую таблицу
        if len(left_table_data) >= 2:
            left_header = clean_headers(left_table_data[0])  # Обрабатываем заголовки
            left_table_data = align_columns(left_table_data[1:],
                                            len(left_header))  # Убираем строку с заголовком из данных
            current_df = pd.DataFrame(left_table_data, columns=left_header)
            current_df = ensure_unique_columns(current_df)  # Гарантируем уникальность заголовков

        # Проверяем, есть ли следующая таблица для возможного объединения
        if i + 1 < num_tables:
            right_table_data = fill_merged_cells(extract_table_data_with_merged_cells(doc.tables[i + 1]))

            # Проверяем, являются ли таблицы частями структуры "регионы слева и справа"
            if is_combined_table(left_table_data, right_table_data):
                print(f"Таблицы {i + 1} и {i + 2} объединены (регионы слева и справа).")

                if len(right_table_data) >= 2:
                    right_header = clean_headers(right_table_data[0])
                    right_table_data = align_columns(right_table_data[1:], len(right_header))
                    df_right = pd.DataFrame(right_table_data, columns=right_header)
                    df_right = ensure_unique_columns(df_right)  # Гарантируем уникальность заголовков

                # Удаляем столбцы с полностью пустыми значениями
                current_df = current_df.dropna(axis=1, how='all')
                df_right = df_right.dropna(axis=1, how='all')

                # Проверяем наличие столбца с названиями регионов в обеих таблицах
                left_region_column = current_df.columns[0]
                right_region_column = df_right.columns[-1]

                # Объединяем таблицы на основе столбца с названиями регионов
                current_df = pd.merge(current_df, df_right, left_on=left_region_column, right_on=right_region_column,
                                      how='inner')
                i += 1  # Пропускаем правую таблицу, так как она уже обработана

        # Сохраняем текущую таблицу в словарь
        header_key = tuple(current_df.columns)
        if header_key in merged_tables:
            merged_tables[header_key] = pd.concat([merged_tables[header_key], current_df], ignore_index=True)
        else:
            merged_tables[header_key] = current_df

        # Переходим к следующей таблице
        current_df = None  # Обнуляем текущую таблицу
        i += 1

    except Exception as e:
        print(f"Ошибка при обработке таблицы {i + 1}: {e}")
        i += 1

# После объединения всех таблиц, проходим по ним и удаляем строки без регионов
for idx, (header_key, df) in enumerate(merged_tables.items(), start=1):
    df = remove_empty_region_rows(df)  # Удаляем строки без данных о регионе
    df = ensure_unique_columns(df)  # Уникализируем заголовки перед сохранением
    csv_file_path = f"ы/5combined_table_{idx}.csv"
    df.to_csv(csv_file_path, index=False)
    print(f"Объединённая таблица сохранена в файл: {csv_file_path}")
